import pdfplumber
import csv
import re
import pandas as pd
from docx import Document
import datetime

def extract_data_from_pdf(pdf_path):
    students = []
    subjects = []
    subjects_extracted = False  # Flag to ensure subjects are only extracted once

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines = text.split("\n")

                # Extract subjects only from the first occurrence of "COURSE NAME"
                if "COURSE NAME" in text and not subjects_extracted:
                    for i, line in enumerate(lines):
                        if "COURSE NAME" in line:
                            course_index = i
                            break

                    for i in range(course_index + 1, len(lines)):
                        if "SGPA" in lines[i]:
                            break
                        subject_match = re.search(r'\d{6}\s+([A-Z &.]+)', lines[i])
                        if subject_match:
                            subjects.append(subject_match.group(1).strip())
                    subjects_extracted = True  # Mark subjects as extracted

                # Extract student data
                student_blocks = re.split(r'\nSEAT NO.:', text)[1:]
                for block in student_blocks:
                    lines = block.split("\n")
                    seat_no = lines[0].split()[0].strip()
                    name_raw = " ".join(lines[0].split()[2:]).strip()

                    if "MOTHER" in name_raw:
                        name = name_raw.split("MOTHER")[0].strip()
                    else:
                        name = name_raw  # Fallback if "MOTHER" is not found

                    student_data = {"Seat No": seat_no, "Name": name}

                    # Extract grades from the "Grd" column
                    extracted_grades = {}
                    for line in lines:
                        match = re.search(r'(\d{6})\s+([A-Z &.]+).*?\s(A\+|O|B\+|B|C|F|P|AC|IC)\s', line)
                        if match:
                            subject_name = match.group(2).strip()
                            grade = match.group(3).strip()
                            extracted_grades[subject_name] = "F" if grade == "F" or grade == "IC" else "P"

                    # Assign extracted grades, checking if subject is missing
                    for subject in subjects:
                        if subject in extracted_grades:
                            student_data[subject] = extracted_grades[subject]
                        else:
                            student_data[subject] = "F" if any(g == "F" for g in extracted_grades.values()) else "P"

                    students.append(student_data)
    return students, subjects


def write_to_csv(students, subjects, output_path):
    headers = ["Seat No", "Name"] + subjects

    with open(output_path, mode='w', newline='', encoding='utf-8') as file:
        writer = csv.DictWriter(file, fieldnames=headers)
        writer.writeheader()

        for student in students:
            writer.writerow(student)

def analyze_results(csv_file):
    # Read the CSV file
    df = pd.read_csv(csv_file)

    # Create a Word document
    doc = Document()
    doc.add_heading('Student Results Analysis', level=1)

    # Dynamically extract subject names (all columns except 'Seat No', 'Name', and 'Total')
    non_subject_columns = ['Seat No', 'Name', 'Total']
    subjects = [col for col in df.columns if col not in non_subject_columns]

    # Perform analysis and write to Word file
    doc.add_heading('Number of students passed in each subject are : ', level=2)
    for subject in subjects:
        pass_count = df[df[subject] == 'P'].shape[0]
        doc.add_paragraph(f"{subject}: {pass_count} students passed")

    # Generate a unique file name with a timestamp
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = f'Result_Analysis_{timestamp}.docx'

    # Save the Word document
    doc.save(output_file)
    print(f"Analysis of the result has been saved to {output_file}, Thank you!")

def main():
    pdf_path = input("Enter the path to the gazette PDF file: ")
    csv_output = "students_results.csv"

    students_data, subject_list = extract_data_from_pdf(pdf_path)
    write_to_csv(students_data, subject_list, csv_output)

    print(f"CSV file '{csv_output}' generated successfully!")
    analyze_results(csv_output)

def run_analysis_from_pdf(pdf_path):
    csv_output = "students_results.csv"
    students_data, subject_list = extract_data_from_pdf(pdf_path)
    write_to_csv(students_data, subject_list, csv_output)
    analyze_results(csv_output)
    return f"Analysis of {pdf_path} complete. Word document generated."
